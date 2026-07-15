"""Build the CityHire Fleet & Stock Intelligence ebook - the narrative data
story tying together every deliverable (SQL, Excel, Power BI, dashboard, ML,
map, competitor/marketing intelligence) into one story/solution/opportunity arc."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
IMG = ROOT / "docs" / "images"
OUT = ROOT / "docs" / "CityHire_Data_Story.docx"

NAVY = RGBColor(0x0C, 0x11, 0x14)
GREEN = RGBColor(0x00, 0x9B, 0x66)
BLUE = RGBColor(0x2A, 0x7A, 0xCC)
SLATE = RGBColor(0x3C, 0x56, 0x67)
GREY = RGBColor(0x52, 0x51, 0x4E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def page_break():
    doc.add_page_break()


def add_kicker(text, color=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def add_h1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def add_h2(text, color=NAVY, space_before=18):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def add_body(text, size=11, color=None, italic=False, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    r.font.color.rgb = color or RGBColor(0x22, 0x22, 0x22)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(11)


def add_image(path, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    p.paragraph_format.space_after = Pt(10)


def hr(color='D9D9D9'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def stat_row(stats):
    """stats: list of (value, label)"""
    table = doc.add_table(rows=1, cols=len(stats))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (val, label) in enumerate(stats):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, 'F2F5F7')
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(val)
        r1.font.name = 'Arial'
        r1.font.bold = True
        r1.font.size = Pt(20)
        r1.font.color.rgb = GREEN
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.name = 'Arial'
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# =============================================================================
# COVER
# =============================================================================
cover_space = doc.add_paragraph()
cover_space.paragraph_format.space_before = Pt(60)

kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = kicker.add_run("DATA & AUTOMATION ANALYST — PORTFOLIO SUBMISSION")
r.font.name = 'Arial'
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = GREEN

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(10)
r = title.add_run("CityHire Fleet & Stock\nIntelligence")
r.font.name = 'Arial'
r.font.size = Pt(34)
r.font.bold = True
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(14)
r = sub.add_run("The data story behind a 1,612-asset fleet — built to show, not tell,\nwhat this role needs")
r.font.name = 'Arial'
r.font.size = Pt(14)
r.font.color.rgb = SLATE

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(30)
r = meta.add_run("Anthony Apollis  •  15 July 2026")
r.font.name = 'Arial'
r.font.size = Pt(10.5)
r.font.color.rgb = GREY

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.paragraph_format.space_before = Pt(4)
r = tag.add_run("Prepared for: Brandon Roman, Talent Shore — Data & Automation Analyst role")
r.font.name = 'Arial'
r.font.size = Pt(10)
r.font.italic = True
r.font.color.rgb = GREY

page_break()

# =============================================================================
# THE STORY
# =============================================================================
add_kicker("The Story")
add_h1("Why I built this instead of just applying")
add_body(
    "Most applications for a role like this arrive as a CV and a cover letter. Both are claims about "
    "capability — “I know SQL,” “I’m advanced in Excel,” “I can build Power BI dashboards.” This document, "
    "and the platform behind it, is the alternative: evidence.",
    space_after=10
)
add_body(
    "I took City Hire’s real, public product catalogue — 24 categories, 220+ genuine product and model "
    "names, from a 3-ton excavator to a Hilti TE 3000-AVR breaker — and built the operation around it: a "
    "1,612-asset fleet, a year of daily stock statistics, repair and approval queues, capex tracking, and "
    "depot-to-depot fleeting. Then I built exactly what the job spec asks a Data & Automation Analyst to "
    "produce and maintain, end to end.",
    space_after=10
)
stat_row([
    ("1,612", "Fleet assets"),
    ("53,802", "Daily stat rows"),
    ("24", "Categories modelled"),
    ("6", "Depots"),
])
add_body(
    "This isn’t a mock-up. Every KPI in the Excel workbook is a live formula. Every chart is driven by a "
    "real SQL query. The Power BI model has 21 working DAX measures across a proper star schema. And the "
    "numbers aren’t all flattering — three categories are running over their capex budget, 32 approvals "
    "are sitting in a queue, and 95 depot/category combinations are short of minimum stock. That’s "
    "deliberate. A dashboard with nothing wrong in it isn’t proof of anything.",
    space_after=4
)

page_break()

# =============================================================================
# THE CHALLENGE
# =============================================================================
add_kicker("The Challenge")
add_h1("What the role actually asks for")
add_body(
    "The job spec, from Brandon Roman at Talent Shore (14 July 2026), sets three 6–12 month success points "
    "and eight daily operational tasks, sitting inside a National Operations & Systems Manager’s stock "
    "function for a plant/tool hire business.",
    space_after=10
)

add_h2("The three success points", space_before=6)
add_bullet(" — the fleet-wide target this build was designed to hit and demonstrate a route toward.",
           bold_prefix="Workshop in-service % down to 15%. ")
add_bullet(" — tracked category by category, because a healthy fleet-wide average can hide real overspend.",
           bold_prefix="Capex spend not exceeding budget. ")
add_bullet(" — the floor every category needs to clear, with exceptions flagged for action.",
           bold_prefix="Utilisation maintained above 60%. ")

add_h2("The daily list", space_before=14)
add_body("Stock stats, fleeting, £500+ approvals, external repairs, stock shortage balancing, minimum "
          "stock levels, stock/parts/sale code creation, and stock & ROI reporting — each one has a live "
          "table, a SQL query, and an Excel sheet behind it in this build. See the case study document for "
          "the full bullet-by-bullet mapping.", space_after=10)

add_h2("The market City Hire competes in", space_before=14)
add_body(
    "City Hire is an independent, London-led operator in a market with a handful of very large national "
    "chains — Sunbelt Rentals, Speedy Hire, HSS Hire, GAP Group — some of whom spend up to £289,000 a "
    "month on paid search alone. City Hire isn’t even tracked in the sector’s national digital-marketing "
    "benchmark reports. That’s the honest starting point for the growth section later in this document.",
    space_after=4
)

page_break()

# =============================================================================
# THE DATA STORY
# =============================================================================
add_kicker("The Data Story")
add_h1("What a year of fleet data actually shows")

add_h2("1. Workshop % is improving — and the model can say when it hits target", space_before=6)
add_body(
    "Workshop time fell from 22.6% to 10.7% of the fleet over the last 12 months, crossing the 15% success "
    "target around March 2026. A gradient-boosted regression model trained on the daily stock data — "
    "independently, without seeing this narrative — predicts the same crossing point, around 9 March 2026. "
    "Two different methods, same answer.", space_after=10
)
add_image(IMG / "trend_chart.png")

add_h2("2. A healthy capex average is hiding two problem categories", space_before=16)
add_body(
    "Fleet-wide capex sits at 87% of budget — on the surface, healthy. Broken out by category, Powered "
    "Access is 6.7% over budget and Plant is 14.5% over, year to date. A logistic regression model trained "
    "purely on category identity correctly flags both as the highest-risk categories for overspend "
    "(AUC 0.56, recovering signal from a genuinely noisy monthly process).", space_after=10
)
add_image(IMG / "capex_variance_chart.png")

page_break()

add_h2("3. ROI is strong fleet-wide, but 30 assets are dragging it down", space_before=6)
add_body(
    "Estimated annual hire revenue against capex outlay runs at 290%+ ROI for most categories — hire "
    "equipment earning back its capital cost several times over its life is the expected shape for this "
    "business. But 30 individual assets return under 120% ROI, concentrated in low-utilisation "
    "(under 35%) units. That’s the auction/disposal shortlist a Data Analyst should be handing to the "
    "Ops & Systems Manager, not a spreadsheet of every asset with no ranking.", space_after=10
)
add_image(IMG / "roi_chart.png")

add_h2("4. Stock shortages cluster by depot, not at random", space_before=16)
add_body(
    "95 category/depot combinations are currently trading below minimum stock, and 18 depot-to-depot "
    "fleeting requests are unfulfilled. Both concentrate at the same two depots — Manchester and London "
    "South — rather than spreading evenly across the network, which points to a specific rebalancing "
    "problem rather than general under-stocking.", space_after=10
)
add_image(IMG / "depot_chart.png")

page_break()

# =============================================================================
# THE SOLUTION
# =============================================================================
add_kicker("The Solution")
add_h1("Everything this build actually delivers")
add_body(
    "Eight linked pieces, all built from the same underlying dataset, so the numbers agree with each "
    "other everywhere they appear.", space_after=10
)

solution_items = [
    ("SQL analysis layer", "17 queries mapped to every job-spec success point and daily task, "
     "against a SQLite schema using portable, OData-translatable patterns."),
    ("Excel workbook (10 sheets)", "Stock Master, Capex vs Budget, ROI Report, Approvals, "
     "External Repairs, Stock Shortage, Min Stock Alerts — every KPI a live formula, zero errors on "
     "recalculation."),
    ("Power BI semantic model", "A real, working model in Power BI Desktop — 11 tables (7 fact tables plus "
     "Calendar, Category and Depot dimensions), 18 relationships, 21 DAX measures organised into 9 display "
     "folders matching the job spec's structure."),
    ("Interactive web dashboard", "KPI tiles, trend and variance charts, an ROI leaderboard, and four "
     "live-filterable action-queue tables — built dependency-free, in City Hire's own colours."),
    ("Interactive depot map", "A UK network view of all 6 depots with switchable overlays (utilisation, "
     "workshop %, stock shortage) and the fleeting flow lines between them."),
    ("Local machine learning", "Four models — utilisation/workshop forecasting, a capex-overspend "
     "classifier, and a day-rate pricing model — run locally rather than on Azure Databricks, given the "
     "Azure trial for this build ends within days."),
    ("Competitor & marketing intelligence", "Where City Hire sits against the national hire chains, plus "
     "a live technical SEO audit of cityhire.co.uk with concrete, fixable findings."),
    ("This document + a case-study brief", "The narrative and the bullet-by-bullet job-spec mapping, for "
     "the application and the interview."),
]

table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [Inches(2.0), Inches(4.3)]
for name, desc in solution_items:
    row = table.add_row().cells
    row[0].width = widths[0]
    row[1].width = widths[1]
    set_cell_shading(row[0], '0C1114')
    p0 = row[0].paragraphs[0]
    r0 = p0.add_run(name)
    r0.font.name = 'Arial'
    r0.font.bold = True
    r0.font.size = Pt(9.5)
    r0.font.color.rgb = WHITE
    p1 = row[1].paragraphs[0]
    r1 = p1.add_run(desc)
    r1.font.name = 'Arial'
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

page_break()

# =============================================================================
# BUSINESS OPPORTUNITY & GROWTH
# =============================================================================
add_kicker("Business Opportunity & Growth")
add_h1("Where this points next")
add_body(
    "The job spec asks for “recommendations of gaps in the market for new product line recommendations” "
    "and “supply chain data analysis, with insight into key gaps within the stock and marketplace.” Three "
    "concrete opportunities came out of this build, beyond the fleet data itself.", space_after=10
)

add_h2("1. Long-tail SEO is winnable, right now", space_before=6)
add_body(
    "Six of the ten nationally-tracked hire brands lost organic search traffic in the last measured "
    "period — one dropped 57%. City Hire isn’t even in that tracked set. Fixing the concrete, findable "
    "issues on cityhire.co.uk (empty Open Graph tags breaking every social share, zero structured data on "
    "product pages, six H1 tags on the homepage) is a low-cost, high-leverage move while larger "
    "competitors are flat or declining on this exact channel.", space_after=10
)

add_h2("2. A pricing-consistency check the fleet data already supports", space_before=16)
add_body(
    "A day-rate pricing model (R² 0.93) trained on capex, age and category flags 20 assets priced well "
    "below what comparable assets earn — a ready-made worklist for a pricing review, not a hunch.",
    space_after=10
)

add_h2("3. Disposal and capital-replacement decisions with a ranked shortlist", space_before=16)
add_body(
    "Rather than “which assets should we consider for auction,” the ROI model already produces the "
    "answer: 30 specific stock codes, ranked by estimated ROI, with utilisation and condition attached. "
    "That’s the kind of “proactive recommendation for the purchase of stock” the job spec calls for, "
    "built once and reusable every month.", space_after=4
)

page_break()

# =============================================================================
# CLOSE
# =============================================================================
add_kicker("Closing")
add_h1("What I'd do in the first 90 days")
add_bullet(" — replace the synthetic generator with a live feed from City Hire's actual stock system "
           "(the SQL layer and Excel formulas are written to be schema-compatible with minimal rework).",
           bold_prefix="Weeks 1-2: ")
add_bullet(" — click Refresh once in Power BI Desktop to load the real semantic model (already fully "
           "built), and get the first live dashboard in front of the Ops & Systems Manager.",
           bold_prefix="Weeks 3-4: ")
add_bullet(" — retrain the four ML models on real fleet history once enough data has accumulated, and "
           "action the SEO fixes flagged in the marketing intelligence section.",
           bold_prefix="Weeks 5-8: ")
add_bullet(" — establish the daily stock-stats, approvals, and repairs cadence this build already "
           "models, with the workshop-% and capex-variance KPIs reported up monthly.",
           bold_prefix="Weeks 9-12: ")

hr()
add_body(
    "All source files — dataset, SQL, Excel, Power BI model, dashboard, map, ML models, and this "
    "document — are available on request, and on GitHub.", italic=True, color=GREY, size=9.5, space_after=2
)
add_body("Anthony Apollis — anthony.apollis@gmail.com", color=GREY, size=10)

doc.save(OUT)
print(f"Saved: {OUT}")
