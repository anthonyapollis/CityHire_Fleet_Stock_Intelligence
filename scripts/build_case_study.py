"""Build the CityHire case-study Word document mapping every deliverable to the
Data & Automation Analyst job spec (Talent Shore / Brandon Roman, 2026-07-14)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "CityHire_Case_Study.docx"

NAVY = RGBColor(0x1F, 0x2A, 0x44)
TEAL = RGBColor(0x0E, 0x7C, 0x7B)
GREY = RGBColor(0x52, 0x51, 0x4E)
CRIT = RGBColor(0xC0, 0x39, 0x2B)
GOOD = RGBColor(0x1E, 0x82, 0x4C)

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

for section in doc.sections:
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(text, level=1, color=NAVY, size=None, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(size or {1: 20, 2: 14, 3: 11.5}.get(level, 11))
    if level == 1:
        p.paragraph_format.space_after = Pt(2)
    return p


def add_body(text, size=10.5, color=None, italic=False, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.bold = bold
    run.font.color.rgb = color or RGBColor(0x22, 0x22, 0x22)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
    return p


def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D9D9D9')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================== COVER ==============================
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(40)
run = title.add_run('CityHire Fleet & Stock Intelligence')
run.font.name = 'Arial'
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(4)
run = sub.add_run('A working case study for the Data & Automation Analyst application')
run.font.name = 'Arial'
run.font.size = Pt(13)
run.font.color.rgb = TEAL

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(2)
run = meta.add_run('Prepared by Anthony Apollis  |  15 July 2026')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.color.rgb = GREY

meta2 = doc.add_paragraph()
meta2.paragraph_format.space_after = Pt(18)
run = meta2.add_run('In response to the role shared by Brandon Roman, Talent Shore (14 July 2026)')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = GREY
hr()

add_body(
    'Rather than send a CV alone, I built the kind of analysis this role actually asks for. '
    'I pulled the real product and category structure from City Hire’s public site, modelled a fleet '
    'and stock operation around it, and built the SQL, Excel and dashboard layer a Data & Automation '
    'Analyst would be expected to produce and maintain — daily stock stats, workshop in-service %, '
    'utilisation, capex vs budget, ROI reporting, and the approval/repair/stock-shortage queues.',
    space_after=10,
)
add_body(
    'IMPORTANT — all fleet quantities, stock codes, costs, dates, utilisation figures, repairs, '
    'approvals and transfers in this build are SYNTHETIC, generated for this demonstration. Only the '
    'category structure and 220+ individual product/model names (e.g. “Nifty Lift HR21 20.8m Hybrid '
    'Boom Lift”, “Hilti TE 3000-AVR Heavy Breaker”) are real, extracted from the public product '
    'catalogue at cityhire.co.uk on 15 July 2026.',
    italic=True, color=GREY, space_after=14,
)

# ============================== WHAT'S IN THE BUILD ==============================
add_heading('What’s in the build', level=2, color=NAVY)
add_body('Four linked deliverables, all built from one 1,612-asset synthetic fleet dataset:', space_after=6)
add_bullet(' — 1,612 assets across 24 categories, 365 days of daily stock stats, external repairs, '
           'auction & parts approvals, capex vs budget, and depot-to-depot transfer logs.',
           bold_prefix='Dataset (CSV + SQLite)')
add_bullet(' — 17 queries covering every success point and daily task in the job spec — workshop '
           'in-service %, utilisation, capex variance, ROI, minimum stock levels, approvals and repairs.',
           bold_prefix='SQL analysis layer')
add_bullet(' — Stock Master, Capex vs Budget, ROI Report, Approvals, External Repairs, Stock Shortage '
           'and Min Stock Alerts — every KPI a live formula (SUMIFS/AVERAGEIFS/COUNTIFS), not a pasted number.',
           bold_prefix='Excel workbook, 10 sheets')
add_bullet(' — executive KPI tiles, a 12-month trend chart, a capex variance chart, an ROI leaderboard '
           'and four live action-queue tables, built as a self-contained interactive page.',
           bold_prefix='Interactive dashboard')

# ============================== JD MAPPING TABLE ==============================
add_heading('How this maps to the job spec', level=2, color=NAVY)
add_body('Every 6–12 month success point and daily task from the spec, paired with the specific '
          'deliverable and artefact that demonstrates it.', space_after=8)

mapping = [
    ("6–12 month: workshop in-service % down to 15%",
     "Monthly Trend sheet / dashboard trend chart",
     "Modelled improving from 22.6% (Jul 2025) to 10.7% (Jul 2026) workshop, crossing the 15% target around March 2026."),
    ("6–12 month: capex spend not exceeding budget",
     "Capex vs Budget sheet (SUMIFS by category) / dashboard variance chart",
     "3 of 24 categories running over budget YTD — Powered Access +£120k (106.7%), Plant +£112k (114.5%), "
     "Survey & Location +£9k (104.2%) — flagged as a capex review list, not hidden by a healthy fleet-wide average."),
    ("6–12 month: utilisation maintained above 60%",
     "SQL query 2 / ROI Report sheet",
     "Fleet-wide utilisation holds a 52–66% seasonal band; only 1 of 24 categories (Fencing & Barriers, 59.6%) sits below the 60% floor."),
    ("Daily: stock stats",
     "daily_stock_stats table (53,802 rows) / SQL query 4",
     "Per-depot, per-category units on hire / in workshop / available, computed daily across the fleet."),
    ("Daily: fleeting",
     "stock_transfers table / Stock Shortage sheet",
     "150 depot-to-depot transfers logged; 18 unfulfilled requests surfaced as an open action queue."),
    ("Daily: auction & parts (£500+) approvals",
     "approvals table / Approvals sheet & dashboard tab",
     "144 requests (auction disposal, parts purchase, capital replacement) with a live pending-approval count (32) and approver routing."),
    ("Daily: external repairs",
     "external_repairs table / External Repairs sheet",
     "90 supplier repair jobs, cost and days-out-of-service tracked per supplier; 18 repairs £500+ awaiting sign-off."),
    ("Daily: stock shortage balancing",
     "SQL query 5 / Stock Shortage sheet",
     "Transfers grouped by category and depot pair, with an unfulfilled-request list for same-day escalation."),
    ("Daily: minimum stock levels",
     "SQL query 6 / Min Stock Alerts sheet",
     "95 category/depot combinations currently trading below their minimum holding, flagged for replenishment."),
    ("Daily: creation of stock/parts/sale codes",
     "sale_codes table",
     "120 hire-stock, parts and sale codes registered against stock items, with status and creator tracked."),
    ("Production of stock and ROI reports",
     "ROI Report sheet / SQL query 9",
     "Estimated annual hire revenue vs capex by category and by individual asset; 30 lowest-ROI assets listed as a disposal/auction review shortlist."),
    ("SQL and OData specialist",
     "sql/analysis_queries.sql",
     "17 queries against a SQLite schema — portable T-SQL patterns (SUMIFS-equivalents, date functions) that map directly onto an OData-backed source."),
    ("Advanced Excel / formula creation",
     "Whole workbook",
     "Every KPI and summary cell (SUMIFS, AVERAGEIFS, COUNTIFS, SUMPRODUCT, IFERROR) is a live formula against the raw data sheets — zero hardcoded results, zero formula errors on recalculation."),
    ("Power BI / technical dashboards",
     "Interactive HTML dashboard",
     "KPI tiles, trend and variance charts, ROI leaderboard and live-filterable action queues — the same information architecture as a Power BI report, built dependency-free."),
    ("Recommendations for new product lines / market gaps",
     "SQL query 10 (undersupplied_categories)",
     "Flags categories sustaining >85% utilisation with low workshop time — the signal for capex expansion or an adjacent product line."),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [Inches(2.1), Inches(1.7), Inches(3.0)]
for i, w in enumerate(widths):
    table.columns[i].width = w

tr = table.rows[0]._tr
trPr = tr.get_or_add_trPr()
tblHeader = OxmlElement('w:tblHeader')
tblHeader.set(qn('w:val'), 'true')
trPr.append(tblHeader)

hdr = table.rows[0].cells
headers = ['Job Spec Item', 'Deliverable', 'What it shows']
for i, h in enumerate(headers):
    hdr[i].width = widths[i]
    set_cell_shading(hdr[i], '1F2A44')
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for i, (spec, deliverable, note) in enumerate(mapping):
    row = table.add_row().cells
    for j, (val, w) in enumerate(zip([spec, deliverable, note], widths)):
        row[j].width = w
        p = row[j].paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.bold = (j == 0)
        r.font.color.rgb = NAVY if j == 0 else RGBColor(0x33, 0x33, 0x33)
        if i % 2 == 1:
            set_cell_shading(row[j], 'F2F2F2')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ============================== KEY FINDINGS ==============================
add_heading('Findings worth a second look', level=2, color=NAVY, space_before=18)
add_body('The point of building this rather than describing it: real numbers surface real problems, '
          'which is the actual job. Four findings the data itself raised:', space_after=8)
add_bullet(' — Powered Access and Plant are running 6.7% and 14.5% over their YTD capex budget '
           'respectively, while the fleet-wide capex position looks healthy at 87% of budget used. '
           'A category-only view hides this; a category-level SUMIFS breakdown catches it.',
           bold_prefix='Capex overspend is hiding inside a healthy average. ')
add_bullet(' — 32 requests over £500 are sitting in the approval queue, several open since January 2026 '
           '— an ageing-approvals report (by days open, not just by value) would be the natural next '
           'build to prioritise sign-off.',
           bold_prefix='Approvals are ageing. ')
add_bullet(' — 30 individual assets return under 120% estimated annual ROI against a fleet median '
           'north of 290% — concentrated in low-utilisation (<35%) units across Material Handling, '
           'Powered Access and Cutting & Grinding. These are the auction/disposal shortlist, not the '
           'lowest-value assets by book value alone.',
           bold_prefix='A small set of assets is dragging fleet ROI. ')
add_bullet(' — 95 category/depot combinations are currently below minimum stock, concentrated at '
           'Manchester and London South — consistent with the 18 unfulfilled fleeting requests, '
           'suggesting a rebalancing pattern rather than isolated one-off shortages.',
           bold_prefix='Stock shortages cluster by depot, not randomly. ')

# ============================== TECHNICAL APPROACH ==============================
add_heading('Technical approach', level=2, color=NAVY, space_before=16)
add_bullet(' — category structure, subcategories and 220+ real product/model names scraped directly '
           'from cityhire.co.uk’s public product pages (7 categories in depth, extrapolated in the same '
           'style for the remaining 17 to cover the full catalogue).',
           bold_prefix='Data grounding: ')
add_bullet(' — Python-generated fleet, stats, repairs, approvals, capex and transfer data at realistic '
           'volume (1,612 assets, 53,802 daily stat rows), with an explicit improvement trend on workshop % '
           'and deliberately-seeded overspend categories so the analysis has real signal to find.',
           bold_prefix='Synthetic data design: ')
add_bullet(' — SQLite schema, 17 analysis queries covering every daily task and success metric in the spec.',
           bold_prefix='SQL layer: ')
add_bullet(' — openpyxl-built, 10 sheets, every KPI a live formula, conditional formatting on utilisation '
           'and status, verified with a full recalculation pass and a zero-formula-error check.',
           bold_prefix='Excel layer: ')
add_bullet(' — self-contained HTML/SVG/JS, no external chart library, hover tooltips, live tab-filtered '
           'action queues, light/dark theme support.',
           bold_prefix='Dashboard layer: ')

hr()
add_body('All files (dataset, SQL, Excel workbook, dashboard, this document) are available on request.',
          italic=True, color=GREY, size=9.5, space_after=2)
add_body('Anthony Apollis  —  anthony.apollis@gmail.com', color=GREY, size=9.5)

doc.save(OUT)
print(f"Saved: {OUT}")
