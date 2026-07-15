"""Build the Competitor & Marketing Intelligence document - the 'beyond the job
spec' add-on: competitive landscape + a real technical SEO pass on cityhire.co.uk."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "CityHire_Competitor_Marketing_Intelligence.docx"

NAVY = RGBColor(0x0C, 0x11, 0x14)
GREEN = RGBColor(0x00, 0x9B, 0x66)
SLATE = RGBColor(0x2D, 0x41, 0x4D)
GREY = RGBColor(0x52, 0x51, 0x4E)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
for section in doc.sections:
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(text, size=14, color=NAVY, space_before=16, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(size)
    return p


def add_body(text, size=10.5, color=None, italic=False, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    r.font.color.rgb = color or RGBColor(0x22, 0x22, 0x22)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)


def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D9D9D9')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================== COVER ==============================
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(30)
r = title.add_run('Competitor & Marketing Intelligence')
r.font.name = 'Arial'
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(16)
r = sub.add_run('Beyond the brief — going past the job spec, addendum to the CityHire Fleet & Stock Intelligence build')
r.font.name = 'Arial'
r.font.size = Pt(12.5)
r.font.color.rgb = GREEN
hr()

add_body(
    'The job spec asks for stock and fleet analytics. Two of the JD bullets — "supply chain data analysis, '
    'with insight into key gaps within the stock and marketplace" and "recommendations of gaps in the market '
    'for new product line recommendations" — reach past the fleet data itself into market positioning. This '
    'section covers that ground: where City Hire sits against the UK’s national hire chains, and a real, '
    'live technical SEO pass on cityhire.co.uk (checked 15 July 2026).', space_after=12
)

# ============================== COMPETITIVE LANDSCAPE ==============================
add_heading('The UK tool & plant hire landscape')
add_body(
    'City Hire operates in a market dominated by a small number of very large national chains, alongside a '
    'long tail of independent and regional hire companies. The scale gap is the defining fact of this market.',
    space_after=8
)

competitors = [
    ("Sunbelt Rentals UK & Ireland (Ashtead Group)", "National, multi-sector",
     "Largest player by group revenue; Ashtead Group's overall scale (employees and revenue) exceeds Speedy Hire's."),
    ("Speedy Hire", "National, 180+ depots",
     "Reaches ~85% of the UK's top contractors; leads UK tool-hire brands on organic search visibility (4,126 universal search appearances, most Instagram followers)."),
    ("HSS Hire", "National, est. 1957",
     "Long-established heavy-equipment specialist; highest measured monthly ad spend in the sector (~£289k) and most Facebook likes (21,500); recently sold £35m of assets/depot leases to Speedy."),
    ("GAP Group", "National, independent",
     "Large independent operator competing on fleet diversification and customer service rather than group scale."),
    ("Brandon Hire Station", "National network",
     "Dominates longtail keyword position-3 rankings (1,979 appearances) — a strong long-tail SEO player despite being a smaller brand than Speedy/HSS."),
    ("City Hire", "Independent, London-led same-day delivery",
     "Not tracked in the sector's national digital-marketing benchmark reports at all — competes on same-day London/regional delivery and personal service, not ad spend or depot count."),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [Inches(2.0), Inches(1.5), Inches(3.3)]
for i, w in enumerate(widths):
    table.columns[i].width = w
tr = table.rows[0]._tr
trPr = tr.get_or_add_trPr()
tblHeader = OxmlElement('w:tblHeader')
tblHeader.set(qn('w:val'), 'true')
trPr.append(tblHeader)
hdr = table.rows[0].cells
for i, h in enumerate(['Competitor', 'Positioning', 'Signal']):
    hdr[i].width = widths[i]
    set_cell_shading(hdr[i], '0C1114')
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for i, (name, pos, note) in enumerate(competitors):
    row = table.add_row().cells
    for j, (val, w) in enumerate(zip([name, pos, note], widths)):
        row[j].width = w
        p = row[j].paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.bold = (j == 0)
        r.font.color.rgb = NAVY if j == 0 else RGBColor(0x33, 0x33, 0x33)
        if name == "City Hire":
            set_cell_shading(row[j], 'E8F5EF')
        elif i % 2 == 1:
            set_cell_shading(row[j], 'F2F5F7')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_heading('Where the market gap actually is', size=13, space_before=16)
add_bullet(' — HSS spends up to ~£289k/month on paid search; Speedy leads organic reach with 4,126 universal '
           'search appearances. City Hire cannot out-spend or out-scale that. The gap that matters isn’t '
           'ad budget, it’s that City Hire isn’t appearing in these sector benchmark reports at all — '
           'meaning its digital footprint isn’t yet being measured by the tools advertisers/analysts use to '
           'judge the category, a visibility gap before it’s a spend gap.',
           bold_prefix='Scale gap, not a service gap. ')
add_bullet(' — Brandon Hire Station’s strength is in long-tail keyword positions (position 3, 1,979 '
           'appearances) rather than head terms — proof that a smaller player can win specific, high-intent '
           'searches (e.g. "5-ton excavator hire Hackney") without matching the majors on broad-term spend. '
           'City Hire’s same-day-delivery, London-first positioning is exactly this kind of long-tail '
           'differentiator, if the technical SEO catches up (see below).',
           bold_prefix='Long-tail search is winnable. ')
add_bullet(' — Six of ten tracked national brands lost desktop organic traffic in the last measured period '
           '(one, Premier Plant Hire, dropped 57%). Structured, consistent SEO/content investment right now is '
           'a genuine opportunity while larger competitors are stagnant or declining on this channel.',
           bold_prefix='The category’s SEO is currently weak, fleet-wide. ')

# ============================== MARKETING INTELLIGENCE / SEO AUDIT ==============================
add_heading('Technical SEO — a live check of cityhire.co.uk', space_before=20)
add_body('Checked directly against the live site on 15 July 2026 (homepage + a product category page). '
          'Real findings, not a generic checklist.', space_after=8)

seo_findings = [
    ("Multiple H1 tags on the homepage", "Critical",
     "The homepage carries 6 <h1> elements (“Live notifications”, “Order Approvals”, "
     "“Custom reports”, “24/7 Online Hire & Pickup Management”, “Testimonials”, "
     "“Supported projects for”) where feature-card headings should be <h2>/<h3>. This dilutes the "
     "page’s single topical signal to search engines."),
    ("Open Graph tags present but empty", "Critical",
     "Category pages (checked: /products/cutting-grinding/) declare og:title, og:type, og:url, og:image and "
     "og:image:alt — all with blank content. Shared links on LinkedIn, Facebook or WhatsApp render with no "
     "title, description or image."),
    ("No structured data on product/category pages", "High",
     "Zero JSON-LD schema blocks found on the category page tested (the homepage carries exactly one). No "
     "LocalBusiness, Product, or Service schema — a missed opportunity given the site’s emphasis on "
     "same-day delivery and multiple depot locations, both strong local-SEO signals when marked up."),
    ("4 of 37 images missing alt text", "Medium",
     "On the tested category page, 4 product images have no alt attribute — an accessibility and image-search issue."),
    ("Robots.txt and sitemap are correctly configured", "Passing",
     "robots.txt blocks only CMS/system paths (Umbraco) and correctly declares the sitemap at "
     "/sitemapxml/. No accidental blocking of product pages found."),
    ("Single, clear H1 and mobile viewport on category pages", "Passing",
     "The tested category page has exactly one H1 (“Hire Cutting, Sawing and Grinding Tools and "
     "Equipment”) and a correct responsive viewport meta tag."),
]

table2 = doc.add_table(rows=1, cols=3)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
widths2 = [Inches(2.2), Inches(0.9), Inches(3.7)]
for i, w in enumerate(widths2):
    table2.columns[i].width = w
tr2 = table2.rows[0]._tr
trPr2 = tr2.get_or_add_trPr()
tblHeader2 = OxmlElement('w:tblHeader')
tblHeader2.set(qn('w:val'), 'true')
trPr2.append(tblHeader2)
hdr2 = table2.rows[0].cells
for i, h in enumerate(['Finding', 'Severity', 'Detail']):
    hdr2[i].width = widths2[i]
    set_cell_shading(hdr2[i], '0C1114')
    p = hdr2[i].paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

sev_colors = {"Critical": RGBColor(0xC0, 0x39, 0x2B), "High": RGBColor(0xD9, 0x8E, 0x04),
              "Medium": RGBColor(0x52, 0x51, 0x4E), "Passing": GREEN}
for i, (finding, sev, detail) in enumerate(seo_findings):
    row = table2.add_row().cells
    for j, val in enumerate([finding, sev, detail]):
        row[j].width = widths2[j]
        p = row[j].paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.bold = (j == 0 or j == 1)
        r.font.color.rgb = sev_colors.get(sev, NAVY) if j == 1 else (NAVY if j == 0 else RGBColor(0x33, 0x33, 0x33))
        if sev == "Passing":
            set_cell_shading(row[j], 'E8F5EF')
        elif i % 2 == 1:
            set_cell_shading(row[j], 'F2F5F7')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_heading('What I’d prioritise first', size=13, space_before=16)
add_bullet(' — highest-leverage single fix: restores social-share previews on every product page, likely a '
           'CMS template fix rather than a content rewrite.', bold_prefix='1. Fix empty Open Graph tags. ')
add_bullet(' — LocalBusiness schema (depot addresses, service areas, opening hours) is the single highest-ROI '
           'addition given the same-day-delivery, multi-depot positioning — it directly targets local-intent '
           'search where the national chains’ broad-match strength matters less.',
           bold_prefix='2. Add structured data. ')
add_bullet(' — demote the feature-card headings to h2/h3 so the page has one clear topical signal.',
           bold_prefix='3. Fix the homepage heading structure. ')
add_bullet(' — Brandon Hire Station’s longtail dominance and the sector-wide organic traffic decline both '
           'point the same way: a content push aimed at long-tail, location-specific search terms is winnable '
           'right now, while the majors are flat or declining on this channel.',
           bold_prefix='4. Invest in long-tail category + location content. ')

hr()
add_body(
    'Sources: sector competitor data and Q1 2026 digital marketing benchmarks via public web research '
    '(Clickthrough Marketing UK Tool & Equipment Hire report, Q1 2026; industry competitor profiles). SEO '
    'findings are a direct, live check against cityhire.co.uk on 15 July 2026, not a template checklist.',
    italic=True, color=GREY, size=9, space_after=2
)
add_body('Anthony Apollis — anthony.apollis@gmail.com', color=GREY, size=9.5)

doc.save(OUT)
print(f"Saved: {OUT}")
