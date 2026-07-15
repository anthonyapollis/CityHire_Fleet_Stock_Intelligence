"""
The single, consolidated CityHire Fleet & Stock Intelligence ebook.
Replaces the three earlier separate documents (case study, market
intelligence, data story) with one document. No recruiter/agency
references - written as a standalone capability demonstration.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
IMG = ROOT / "docs" / "images"
PROD = IMG / "products"
OUT = ROOT / "docs" / "CityHire_Fleet_Intelligence_Ebook.docx"

NAVY = RGBColor(0x0C, 0x11, 0x14)
GREEN = RGBColor(0x00, 0x9B, 0x66)
GREEN_DARK = RGBColor(0x00, 0x75, 0x4D)
BLUE = RGBColor(0x2A, 0x7A, 0xCC)
SLATE = RGBColor(0x3C, 0x56, 0x67)
GREY = RGBColor(0x52, 0x51, 0x4E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xC0, 0x39, 0x2B)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def page_break():
    doc.add_page_break()


def add_kicker(text, color=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def add_h1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def add_h2(text, color=NAVY, space_before=18):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def add_body(text, size=11, color=None, italic=False, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    r.font.color.rgb = color or RGBColor(0x22, 0x22, 0x22)
    return p


def add_use_note(text):
    """A short, visually distinct 'how to use this' callout under a chart."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    r1 = p.add_run("USE  ")
    r1.font.name = 'Arial'
    r1.font.size = Pt(9)
    r1.font.bold = True
    r1.font.color.rgb = GREEN_DARK
    r2 = p.add_run(text)
    r2.font.name = 'Arial'
    r2.font.size = Pt(9.5)
    r2.font.italic = True
    r2.font.color.rgb = SLATE
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(11)


def add_image(path, width=6.3, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    p.paragraph_format.space_after = Pt(4 if caption else 10)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.font.name = 'Arial'
        cr.font.size = Pt(8.5)
        cr.font.italic = True
        cr.font.color.rgb = GREY
        cp.paragraph_format.space_after = Pt(10)


def hr(color='D9D9D9'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def stat_row(stats):
    table = doc.add_table(rows=1, cols=len(stats))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (val, label) in enumerate(stats):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, 'F2F5F7')
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(val)
        r1.font.name = 'Arial'
        r1.font.bold = True
        r1.font.size = Pt(20)
        r1.font.color.rgb = GREEN
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.name = 'Arial'
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def styled_table(headers, rows, widths, header_bg='0C1114'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(widths):
        table.columns[i].width = w
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].width = widths[i]
        set_cell_shading(hdr[i], header_bg)
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'Arial'
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = WHITE
    for i, row in enumerate(rows):
        rc = table.add_row().cells
        for j, val in enumerate(row):
            rc[j].width = widths[j]
            p = rc[j].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            r.font.bold = (j == 0)
            r.font.color.rgb = NAVY if j == 0 else RGBColor(0x33, 0x33, 0x33)
            if i % 2 == 1:
                set_cell_shading(rc[j], 'F2F5F7')
    return table


# =============================================================================
# COVER
# =============================================================================
doc.add_paragraph().paragraph_format.space_before = Pt(50)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = kicker.add_run("APPLIED DATA & AUTOMATION — A FLEET INTELLIGENCE PLATFORM")
r.font.name = 'Arial'
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = GREEN

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(10)
r = title.add_run("CityHire Fleet\n& Stock Intelligence")
r.font.name = 'Arial'
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(14)
r = sub.add_run("SQL. Excel. Power BI. A live Azure database. Network graphs, funnels, seasonal\n"
                "forecasting and a recommendation engine — built on a 1,612-asset fleet.")
r.font.name = 'Arial'
r.font.size = Pt(13.5)
r.font.color.rgb = SLATE

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(26)
r = meta.add_run("Anthony Apollis  ·  Data & Automation Analyst")
r.font.name = 'Arial'
r.font.size = Pt(11)
r.font.color.rgb = GREY

if (PROD / "excavator_20ton.jpg").exists():
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    add_image(PROD / "excavator_20ton.jpg", width=3.6)

page_break()

# =============================================================================
# THE PLATFORM AT A GLANCE
# =============================================================================
add_kicker("The Platform")
add_h1("Nine analytics layers, one fleet")
add_body(
    "This platform is a full-stack demonstration of applied data and automation work: real product "
    "data grounding a synthetic fleet, then nine layers of analysis built on top of it — SQL, Excel, "
    "Power BI, a live cloud database, an interactive dashboard, network graphs, funnels, seasonal "
    "forecasting, and a recommendation engine. Every number in every layer traces back to the same "
    "1,612-asset fleet, so nothing here is a disconnected demo chart.",
    space_after=10
)
stat_row([
    ("1,612", "Fleet assets"), ("53,802", "Daily stat rows"), ("5,000", "Synthetic hire orders"),
    ("9", "Analytics layers"), ("6", "UK depots"),
])
add_body(
    "Category structure and 220+ individual product/model names — from a 3-ton excavator to a Hilti "
    "TE 3000-AVR breaker — are real, extracted from City Hire's public product catalogue. Fleet "
    "quantities, costs, utilisation, repairs, approvals, orders and transfers are synthetic, generated "
    "to give every layer below genuine, non-trivial signal to find.",
    italic=True, color=GREY, size=9.5, space_after=4
)

page_break()

# =============================================================================
# SECTION: THE FLEET DATA STORY
# =============================================================================
add_kicker("Fleet Performance")
add_h1("The core numbers, live")
add_body(
    "Three metrics drive fleet operations: how much of the fleet is earning (utilisation), how much is "
    "off the road (workshop time), and whether capital spend is under control (capex vs budget). All "
    "three trace to daily stock-count data across the full 12-month history.",
    space_after=10
)
add_image(IMG / "trend_chart.png")
add_use_note("Track the two headline KPIs against their targets month by month — the gap between "
              "utilisation and 60%, and between in-service % and 85%, is the number to report upward "
              "each period.")

add_image(IMG / "capex_variance_chart.png")
add_use_note("A category-level breakdown catches what a fleet-wide capex average hides — here, Powered "
              "Access and Plant are the two categories actually over budget, everything else is under.")

if (PROD / "telehandler_14m.png").exists():
    add_image(PROD / "telehandler_14m.png", width=2.6,
              caption="14m Telehandler — one of the Plant category's highest-capex assets in the fleet")

page_break()

add_image(IMG / "roi_chart.png")
add_use_note("Rank categories by estimated ROI to prioritise the next capital purchase toward the "
              "categories already paying back fastest.")

add_image(IMG / "depot_chart.png")
add_use_note("Compare depots side by side to spot which site needs operational attention — Bristol and "
              "Manchester run the highest workshop % here, worth a maintenance-scheduling review.")

page_break()

# =============================================================================
# SECTION: SEASONAL INTELLIGENCE
# =============================================================================
add_kicker("Seasonal Intelligence")
add_h1("When demand actually moves, and why")
add_body(
    "A classical time-series decomposition splits fleet utilisation into trend, seasonal cycle and "
    "residual noise — the standard technique for separating \"is this really seasonal\" from random "
    "day-to-day variation.",
    space_after=10
)
add_image(IMG / "seasonal_decomposition.png", width=6.0)
add_use_note("Confirms the seasonal effect is real and roughly 30-day cyclical, not noise — the "
              "trend line is what to report as the underlying direction of travel.")

page_break()

add_image(IMG / "seasonal_demand_index.png")
add_use_note("Read each project type's peak month directly off its line, then check it against current "
              "stock levels ahead of time — Landscaping peaking in April is a February purchasing "
              "decision, not an April one.")

# =============================================================================
# SECTION: NETWORK INTELLIGENCE (GRAPH ANALYSIS)
# =============================================================================
add_h2("Network intelligence — graph analysis", space_before=22)
add_body(
    "Two graph models, built with NetworkX: which categories get hired together (a cross-sell network), "
    "and how stock actually moves between depots (a fleeting network). Centrality and community-detection "
    "algorithms turn raw transaction/transfer data into a structural map of the business.",
    space_after=10
)
add_image(IMG / "category_graph.png")

page_break()
add_image(IMG / "depot_graph.png")

# =============================================================================
# SECTION: FUNNEL INTELLIGENCE
# =============================================================================
add_h2("Funnel intelligence", space_before=22)
add_body(
    "Three funnels, each answering a different operational question: how many quotes actually become "
    "completed hires, how quickly approvals clear, and (directionally) how the digital side of the "
    "business converts.",
    space_after=10
)
add_image(IMG / "order_funnel.png")
add_use_note("The biggest single drop is quote-to-approval (86% carried through) — the stage worth "
              "instrumenting first if conversion needs to improve.")

page_break()
add_image(IMG / "approvals_funnel.png")
add_use_note("A live approval-rate benchmark (72.9% here) — track this monthly to catch a slowing "
              "approvals process before it becomes a bottleneck.")

add_image(IMG / "digital_funnel.png")
add_use_note("No real analytics exist for cityhire.co.uk yet — this shape is what a working GA4 "
              "instrumentation would need to measure for real, not a finding to act on directly.")

page_break()

# =============================================================================
# SECTION: RECOMMENDATION ENGINE
# =============================================================================
add_kicker("Recommendation Engine")
add_h1("What to sell, and what to buy, next")
add_body(
    "Market-basket association rules (support, confidence, lift) over 5,000 orders, the same technique "
    "behind \"customers who bought X also bought Y\" — applied here to hire equipment rather than retail.",
    space_after=10
)
if (PROD / "scissor_lift_16m.png").exists():
    add_image(PROD / "scissor_lift_16m.png", width=2.4,
              caption="16m Electric Scissor Lift — Powered Access, the category with the largest capex "
                      "footprint in the fleet")

assoc_rows = [
    ("Propping & Support", "Cutting & Grinding", "57.3%", "7.5x"),
    ("Fixing & Fastening", "Mechanical & Electrical", "74.9%", "7.4x"),
    ("Safety, Ventilation & Extraction", "Cutting & Grinding", "55.6%", "7.3x"),
    ("Lighting", "Accommodation, Storage & Welfare", "59.2%", "7.1x"),
]
styled_table(["If hired…", "→ also recommend…", "Confidence", "Lift"], assoc_rows,
             [Inches(2.1), Inches(2.3), Inches(1.1), Inches(0.9)])
add_use_note("A lift above 1.0 means the pairing happens more than chance would predict — feed the "
              "top few rows straight into a quote-screen cross-sell prompt.")
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_h2("Proactive stock purchase recommendations", space_before=16)
add_body(
    "Three signals combined into one worklist: which project types are entering their seasonal peak, "
    "which of their categories are already short of minimum stock, and what their historical ROI has "
    "been — directly answering \"what should we buy, and when.\"",
    space_after=8
)
purchase_rows = [
    ("Access & Height", "0 months (now)", "5 depots short", "304%"),
    ("Finishing", "3 months", "6 depots short", "305%"),
    ("Fixing", "3 months", "6 depots short", "305%"),
    ("Demolition", "4 months", "5 depots short", "291%"),
]
styled_table(["Project type", "Time to seasonal peak", "Current shortage", "Category ROI"],
             purchase_rows, [Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.3)])
add_use_note("Sort by time-to-peak — Access & Height is short of stock right now, at its seasonal peak "
              "already, which makes it the most urgent line on this list.")

page_break()

# =============================================================================
# SECTION: COMPETITIVE & MARKET POSITION
# =============================================================================
add_kicker("Market Position")
add_h1("Where City Hire sits in the UK hire market")
if (PROD / "dumper_3ton.jpeg").exists():
    add_image(PROD / "dumper_3ton.jpeg", width=2.6,
              caption="3 Ton Straight Dumper — one of the 220+ real product/model names grounding this "
                      "platform's category structure")
add_body(
    "City Hire competes against a small number of very large national chains — Sunbelt Rentals, Speedy "
    "Hire, HSS Hire, GAP Group — some spending up to £289,000/month on paid search alone. The honest "
    "read: City Hire isn't tracked in the sector's national digital-marketing benchmark reports at all.",
    space_after=10
)
add_bullet(" — six of ten tracked national hire brands lost organic search traffic in the last "
           "measured period (one dropped 57%) — a real opening while larger players are flat or declining.",
           bold_prefix="The category's SEO is currently weak, fleet-wide. ")
add_bullet(" — Brandon Hire Station's strength is long-tail keyword positions, not head terms, proof "
           "a smaller player can win specific, high-intent searches without matching majors on ad spend.",
           bold_prefix="Long-tail search is winnable. ")

add_h2("A live technical SEO check of cityhire.co.uk", space_before=18)
seo_rows = [
    ("Multiple H1 tags on homepage", "Critical", "6 <h1> elements where feature cards should be <h2>/<h3>"),
    ("Open Graph tags present but empty", "Critical", "Every product page shares with a blank preview"),
    ("No structured data on product pages", "High", "Zero JSON-LD - a missed local-SEO opportunity"),
    ("4 of 37 images missing alt text", "Medium", "Accessibility and image-search gap"),
]
styled_table(["Finding", "Severity", "Detail"], seo_rows, [Inches(2.1), Inches(0.9), Inches(3.4)])
add_use_note("Fix Open Graph tags first — single template change, restores every shared-link preview "
              "site-wide.")

page_break()

# =============================================================================
# SECTION: THE TECHNICAL BUILD (including live Azure DB)
# =============================================================================
add_kicker("The Technical Build")
add_h1("How it's actually built")
add_body("Nine layers, each doing a different job, all reading from the same fleet dataset:", space_after=10)

solution_items = [
    ("SQL analysis layer", "17 queries against a SQLite schema, portable to OData - covers every stock/fleet metric a National Ops function tracks daily."),
    ("Live Azure database", "Fleet data loaded live into Azure Database for PostgreSQL (Flexible Server) and queried directly against the cloud instance, not a local stand-in."),
    ("Excel workbook, 10 sheets", "Every KPI a live formula (SUMIFS/AVERAGEIFS/COUNTIFS/SUMPRODUCT) - zero hardcoded results, zero errors on recalculation."),
    ("Power BI semantic model", "Built live in Power BI Desktop via its modeling API - 11 tables, 18 relationships, 21 DAX measures across 9 display folders."),
    ("Interactive dashboard & map", "Self-contained, dependency-free HTML/SVG - KPI tiles, trend/variance charts, an ROI leaderboard, a UK depot network map."),
    ("Local machine learning", "Four models (utilisation/workshop forecasting, capex-overspend classification, day-rate pricing) - deliberately local, not cloud, given the Azure trial window."),
    ("Graph, funnel & recommendation layers", "NetworkX network analysis, funnel conversion tracking, and market-basket recommendations - all detailed above."),
]
styled_table(["Layer", "What it is"], [(a, b) for a, b in solution_items],
             [Inches(1.9), Inches(4.5)])

doc.add_paragraph().paragraph_format.space_after = Pt(10)
add_h2("The live Azure database, in detail", space_before=16)
add_body(
    "Rather than claim cloud-database experience, this build used one: an Azure Database for PostgreSQL "
    "Flexible Server, provisioned, loaded with 55,414 rows across two tables, and queried directly.",
    space_after=8
)
add_bullet(" — SELECT category, SUM(capex_gbp), AVG(utilisation_pct_ytd), and an ROI calculation, "
           "grouped and ranked, returned in under a second against the live instance.",
           bold_prefix="Query executed: ")
add_bullet(" — Powered Access led on estimated annual revenue (£13.8m) among all categories, matching "
           "the local SQLite analysis exactly - the same numbers, a different engine.",
           bold_prefix="Result: ")
add_bullet(" — a subscription-level policy blocks Azure ML, Databricks, Azure SQL Server, VMs, and "
           "Cosmos DB on this trial account. PostgreSQL Flexible Server is not on that list, so it's "
           "what got used - the resource group was deleted immediately after capturing results, so "
           "nothing was left running or billing.",
           bold_prefix="A genuine constraint, handled: ")

page_break()

# =============================================================================
# CLOSE
# =============================================================================
add_kicker("Summary")
add_h1("What this demonstrates")
add_body(
    "SQL, Excel, Power BI, a live cloud database, network analysis, funnel tracking, seasonal "
    "forecasting, and a working recommendation engine - applied to one coherent fleet, not nine "
    "disconnected tutorials. Every number agrees with every other number, because they all come from "
    "the same source.",
    space_after=10
)
add_body(
    "The synthetic dataset was deliberately built with real signal to find, not sanitised for a good "
    "demo: three categories over capex budget, 32 approvals in queue, 95 depot/category combinations "
    "short of stock, 30 assets dragging down fleet ROI. Finding and ranking those is the actual job.",
    space_after=10
)

hr()
add_body("Source files - dataset, SQL, Excel, Power BI model, dashboard, map, and ML models - available on request.",
          italic=True, color=GREY, size=9.5, space_after=2)
add_body("Anthony Apollis — anthony.apollis@gmail.com", color=GREY, size=10)

doc.save(OUT)
print(f"Saved: {OUT}")
